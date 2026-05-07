"""Substitute placeholder values into template source text.

Supports three input forms:
- Markdown / plain text — pure ``str.replace`` after sorting keys longest-first
  so ``[NUMBER OF DAYS]`` substitutes before ``[NUMBER]``.
- DOCX — uses ``python-docx`` if available. DOCX splits text across multiple
  ``runs`` whenever formatting changes, so a placeholder like
  ``[TENANT FULL NAME]`` may be split as ``[TENANT``, ``FULL``, ``NAME]``
  across three runs. We merge consecutive runs in each paragraph before
  substituting, then preserve the merged formatting (font of the first run).
- PDF output (rendered_original) — generated via ``reportlab`` from the
  rendered Markdown. We avoid weasyprint due to its native dependency
  footprint on Windows (per the spec's stop conditions).

When ``python-docx`` is not installed (e.g. in CI without the optional dep),
``render_docx_bytes`` falls back to MD rendering of the extracted text — the
caller is told the fallback fired so it can record the limitation.

Blank-line substitution rules:
- Signature placeholders (any bracketed key matching ``*SIGNATURE``) get a
  blank underscore line so the rendered doc shows a signing line in place of
  literal ``[LANDLORD SIGNATURE]`` text. Signatures are applied at signing
  time, not generation time.
- The bare ``[DATE]`` placeholder (as distinct from ``[EFFECTIVE DATE]``)
  appears next to signature lines ("Date: ____"). It must be left blank for
  the physical signer to write in — the same rule as SIGNATURE keys.
  ``default_source_map`` intentionally has no ``default_source`` for ``DATE``
  so the resolver does not fill it at generate time; this renderer substitutes
  any unfilled ``[DATE]`` with a blank underscore line.
"""
from __future__ import annotations

import io
import logging
import re

logger = logging.getLogger(__name__)

SIGNATURE_LINE = "_______________________________"
_SIGNATURE_KEY_RE = re.compile(r"\[([A-Z][A-Z0-9 _\-]*?SIGNATURE)\]")
# ``[DATE]`` (bare) next to a signature line must be left blank for the signer
# to write in — same treatment as SIGNATURE keys. ``[EFFECTIVE DATE]`` is a
# document property filled at generation time and is NOT matched here.
_DATE_KEY_RE = re.compile(r"\[DATE\]")


def _augment_with_signature_lines(
    template_text: str, values: dict[str, str],
) -> dict[str, str]:
    """Return ``values`` with any unfilled blank-line placeholders substituted.

    Covers two classes of placeholder that must be left blank for the physical
    signer to fill in:

    - ``*SIGNATURE`` keys (e.g. ``[LANDLORD SIGNATURE]``, ``[TENANT SIGNATURE]``)
    - The bare ``[DATE]`` key (date-of-signing, distinct from ``[EFFECTIVE DATE]``)

    Scans ``template_text`` for matching keys not already in ``values`` and
    maps them to the underscore blank line. The caller's ``values`` dict is
    not mutated.
    """
    augmented = dict(values)
    for match in _SIGNATURE_KEY_RE.finditer(template_text):
        key = match.group(1)
        if key not in augmented:
            augmented[key] = SIGNATURE_LINE
    if _DATE_KEY_RE.search(template_text) and "DATE" not in augmented:
        augmented["DATE"] = SIGNATURE_LINE
    return augmented


def _build_substitution_pattern(values: dict[str, str]) -> list[tuple[str, str]]:
    """Return ``[(bracketed_key, replacement_value)]`` sorted longest-key-first.

    Sorting longest-first prevents partial overlap: ``[NUMBER OF DAYS]`` must
    be replaced before ``[NUMBER]`` or the latter would gobble the prefix.
    """
    items = [(f"[{key}]", str(value if value is not None else "")) for key, value in values.items()]
    items.sort(key=lambda kv: len(kv[0]), reverse=True)
    return items


def render_md(template_text: str, values: dict[str, str]) -> str:
    """Replace ``[KEY]`` tokens in ``template_text`` with ``values[KEY]``.

    Pure function. Unknown keys (placeholders not in ``values``) are left as
    bracketed text so the host can spot them in the rendered output. Any
    ``*SIGNATURE`` placeholder absent from ``values`` is replaced with a
    blank signature line.
    """
    augmented = _augment_with_signature_lines(template_text, values)
    output = template_text
    for needle, replacement in _build_substitution_pattern(augmented):
        output = output.replace(needle, replacement)
    return output


def render_docx_bytes(
    docx_bytes: bytes,
    values: dict[str, str],
) -> tuple[bytes, bool]:
    """Render a DOCX file's bytes with placeholder substitutions.

    Returns ``(rendered_bytes, used_docx_library)``. When ``python-docx`` is
    not installed, ``used_docx_library`` is False and the original bytes are
    returned unchanged — the caller should fall back to a generated MD file.
    """
    try:
        import docx  # type: ignore[import-untyped]
    except ImportError:
        logger.warning(
            "python-docx not installed — DOCX rendering disabled, returning original bytes",
        )
        return docx_bytes, False

    try:
        document = docx.Document(io.BytesIO(docx_bytes))
    except Exception:  # noqa: BLE001
        logger.warning("Failed to parse DOCX bytes — returning original", exc_info=True)
        return docx_bytes, False

    # Walk every paragraph (top-level + table cells) to find SIGNATURE keys
    # the caller didn't supply, so the rendered doc shows a blank signing line.
    # Known limitation: section headers/footers and tables nested inside table
    # cells are not walked. The same paragraphs are also untouched by
    # ``_substitute_in_paragraphs`` below, so this mirrors the existing
    # substitution scope rather than expanding it.
    all_text_chunks: list[str] = []
    for paragraph in document.paragraphs:
        all_text_chunks.append("".join(run.text for run in paragraph.runs))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    all_text_chunks.append("".join(run.text for run in paragraph.runs))
    augmented = _augment_with_signature_lines("\n".join(all_text_chunks), values)

    pattern = _build_substitution_pattern(augmented)
    _substitute_in_paragraphs(document.paragraphs, pattern)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                _substitute_in_paragraphs(cell.paragraphs, pattern)

    out = io.BytesIO()
    document.save(out)
    return out.getvalue(), True


def _substitute_in_paragraphs(paragraphs, pattern: list[tuple[str, str]]) -> None:
    """Merge runs in each paragraph, then substitute. ``pattern`` is longest-first.

    The merge step is necessary because Word splits a single visual placeholder
    across multiple runs whenever formatting changes mid-token. If we don't
    merge, ``str.replace`` will not find the bracketed key.
    """
    for paragraph in paragraphs:
        runs = paragraph.runs
        if not runs:
            continue
        # Merge: concatenate all run text into the first run, blank the rest.
        merged_text = "".join(run.text for run in runs)
        for needle, replacement in pattern:
            merged_text = merged_text.replace(needle, replacement)
        runs[0].text = merged_text
        for run in runs[1:]:
            run.text = ""


def render_docx_bytes_to_pdf(
    docx_bytes: bytes, values: dict[str, str],
) -> tuple[bytes, bool]:
    """Render a DOCX template into a PDF after placeholder substitution.

    Pipeline: ``docx_bytes`` → python-docx (substitute placeholders, merge
    runs) → LibreOffice headless (`soffice --convert-to pdf`) for a
    1:1-fidelity PDF that mirrors the source DOCX's layout, fonts,
    tables, and headings.

    Returns ``(pdf_bytes, used_docx_library)``. Falls back to a plain-
    text PDF render via reportlab when LibreOffice is unavailable
    (typically: pytest on a dev machine without the soffice binary). In
    production the docker image installs ``libreoffice-writer`` so the
    high-fidelity branch is the live path.
    """
    import shutil
    import subprocess
    import tempfile

    rendered_docx, used_docx = render_docx_bytes(docx_bytes, values)
    if not used_docx:
        # python-docx not installed; we never substituted. Best we can do
        # is render the raw bytes' decoded text.
        text = docx_bytes.decode("utf-8", errors="replace")
        substituted = render_md(text, values)
        return render_pdf_from_text(substituted), False

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice is None:
        logger.warning(
            "LibreOffice (soffice) not found on PATH — falling back to "
            "plain-text PDF. Install libreoffice-writer for full fidelity.",
        )
        try:
            import docx  # type: ignore[import-untyped]
            doc = docx.Document(io.BytesIO(rendered_docx))
            extracted = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return render_pdf_from_text(extracted), True
        except Exception:  # noqa: BLE001
            return render_pdf_from_text(""), True

    with tempfile.TemporaryDirectory(prefix="lease-pdf-") as tmpdir:
        docx_path = f"{tmpdir}/lease.docx"
        with open(docx_path, "wb") as fh:
            fh.write(rendered_docx)
        # ``--headless`` runs without an X server. ``--convert-to pdf``
        # writes ``lease.pdf`` to ``--outdir``. We use a fresh
        # ``-env:UserInstallation`` per call so concurrent renders don't
        # contend on a shared profile lock (LibreOffice serializes
        # otherwise).
        try:
            subprocess.run(
                [
                    soffice,
                    "--headless",
                    f"-env:UserInstallation=file://{tmpdir}/profile",
                    "--convert-to", "pdf",
                    "--outdir", tmpdir,
                    docx_path,
                ],
                check=True,
                capture_output=True,
                timeout=60,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            logger.warning(
                "soffice DOCX → PDF conversion failed (%s) — falling back "
                "to plain-text PDF.", exc,
            )
            try:
                import docx  # type: ignore[import-untyped]
                doc = docx.Document(io.BytesIO(rendered_docx))
                extracted = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
                return render_pdf_from_text(extracted), True
            except Exception:  # noqa: BLE001
                return render_pdf_from_text(""), True

        with open(f"{tmpdir}/lease.pdf", "rb") as fh:
            pdf_bytes = fh.read()
    return pdf_bytes, True


def render_pdf_from_text(rendered_text: str) -> bytes:
    """Generate a simple PDF from rendered plain/markdown text via ``reportlab``.

    This is intentionally low-fidelity for Phase 1 — it lays out paragraphs
    on letter-size pages with a default font and no markdown styling. The
    fidelity story improves in Phase 2 once we're confident enough to add
    weasyprint or pandoc as a dep.
    """
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )
    except ImportError:  # pragma: no cover — reportlab is in pyproject deps
        raise RuntimeError("reportlab is required for PDF rendering") from None

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=LETTER)
    styles = getSampleStyleSheet()
    body_style = styles["BodyText"]

    story = []
    for raw_paragraph in rendered_text.split("\n\n"):
        text = raw_paragraph.strip()
        if not text:
            story.append(Spacer(1, 12))
            continue
        # Preserve single newlines as <br/> within a paragraph for reportlab.
        safe = (
            text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        ).replace("\n", "<br/>")
        story.append(Paragraph(safe, body_style))
        story.append(Spacer(1, 8))

    doc.build(story)
    return buffer.getvalue()
