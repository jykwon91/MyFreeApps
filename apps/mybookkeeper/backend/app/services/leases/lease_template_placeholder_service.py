"""AI placeholder suggestion for lease templates.

Extracts text from uploaded files and calls Claude to propose a named +
typed list of placeholders that the host can review before persisting.

Constants ``ALLOWED_TEMPLATE_MIME_TYPES`` and ``DOCX_MIME`` are defined here
because both CRUD and placeholder-extraction need them. The CRUD module
imports them from here to keep a single source of truth.
"""
from __future__ import annotations

import logging
import uuid

from app.core.storage import get_storage
from app.db.session import unit_of_work
from app.repositories.leases import (
    lease_template_file_repo,
    lease_template_repo,
)
from app.schemas.leases.suggest_placeholders_response import (
    SuggestedPlaceholderItem,
    SuggestPlaceholdersResponse,
)
from app.services.leases.template_placeholder_extractor import (
    suggest_placeholders as _ai_suggest_placeholders,
)

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# MIME / content-type constants
# ---------------------------------------------------------------------------

ALLOWED_TEMPLATE_MIME_TYPES: frozenset[str] = frozenset({
    "text/markdown",
    "text/plain",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
})

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# ---------------------------------------------------------------------------
# Exceptions (shared with the CRUD module — re-exported from there)
# ---------------------------------------------------------------------------

class TemplateNotFoundError(LookupError):
    pass


class StorageNotConfiguredError(RuntimeError):
    pass


# ---------------------------------------------------------------------------
# Helpers: content-type normalisation + text extraction
# ---------------------------------------------------------------------------

def normalise_content_type(declared: str | None, filename: str) -> str:
    """Resolve a content type when the upload omits one or sends octet-stream.

    DOCX uploads frequently arrive as ``application/octet-stream`` from
    browser drag-drop, so we fall back to the filename extension.
    """
    if declared and declared in ALLOWED_TEMPLATE_MIME_TYPES:
        return declared
    lower = filename.lower()
    if lower.endswith(".docx"):
        return DOCX_MIME
    if lower.endswith(".md"):
        return "text/markdown"
    if lower.endswith(".txt"):
        return "text/plain"
    if declared:
        return declared
    return "application/octet-stream"


def extract_text_from_upload(content: bytes, content_type: str) -> str:
    """Pull plain text out of the upload for placeholder extraction.

    For DOCX we use python-docx if available; otherwise we can't extract
    placeholders (the host can still edit the spec by hand later).
    """
    if content_type in ("text/markdown", "text/plain"):
        try:
            return content.decode("utf-8", errors="replace")
        except Exception:  # noqa: BLE001
            return ""
    if content_type == DOCX_MIME:
        try:
            import io
            import docx  # type: ignore[import-untyped]

            document = docx.Document(io.BytesIO(content))
            parts = [p.text for p in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            parts.append(p.text)
            return "\n".join(parts)
        except ImportError:
            logger.info(
                "python-docx not installed — DOCX placeholder extraction skipped",
            )
            return ""
        except Exception:  # noqa: BLE001
            logger.warning("Failed to extract text from DOCX", exc_info=True)
            return ""
    return ""


# ---------------------------------------------------------------------------
# AI placeholder suggestion
# ---------------------------------------------------------------------------

async def suggest_ai_placeholders(
    *,
    user_id: uuid.UUID,
    organization_id: uuid.UUID,
    template_id: uuid.UUID,
) -> SuggestPlaceholdersResponse:
    """Run an AI pass over the template's files and propose placeholders.

    Fetches file content from MinIO, extracts text from each file, then calls
    Claude to propose a named + typed list of placeholders. The caller
    (route handler) must NOT persist the result — the host reviews the list
    and saves via the existing ``update_placeholder`` endpoint.

    Returns a :class:`SuggestPlaceholdersResponse` with ``suggestions``,
    ``truncated``, and an optional ``pages_note``.

    Raises :class:`TemplateNotFoundError` if the template doesn't belong to the
    caller. Storage errors are propagated (the route handler maps them to 503).
    """
    storage = get_storage()
    if storage is None:
        raise StorageNotConfiguredError("Object storage is not configured")

    async with unit_of_work() as db:
        template = await lease_template_repo.get(
            db,
            template_id=template_id,
            user_id=user_id,
            organization_id=organization_id,
        )
        if template is None:
            raise TemplateNotFoundError(f"Template {template_id} not found")
        files = await lease_template_file_repo.list_for_template(
            db, template_id=template_id,
        )

    extracted_texts: list[str] = []
    for f in files:
        raw = storage.download_file(f.storage_key)
        extracted_texts.append(extract_text_from_upload(raw, f.content_type))

    combined_text = "\n\n".join(t for t in extracted_texts if t.strip())
    result = await _ai_suggest_placeholders(combined_text)

    pages_note: str | None = None
    if result.truncated:
        pages_note = (
            "The document was too long to analyse in full — "
            "I read the first portion. Some placeholders near the end may be missing."
        )

    return SuggestPlaceholdersResponse(
        suggestions=[
            SuggestedPlaceholderItem(
                key=s.key,
                description=s.description,
                input_type=s.input_type,
            )
            for s in result.suggestions
        ],
        truncated=result.truncated,
        pages_note=pages_note,
    )
