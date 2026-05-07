"""Service layer for lease templates.

Per layered-architecture: routes thin, services orchestrate, repositories own
queries. Tenant isolation is via ``(user_id, organization_id)`` per the
project's existing convention. Tenant scope is checked first on every call —
cross-tenant access returns the same response shape as a missing row.

Pipeline for ``upload_template``:
    parse name/desc → upload each file to MinIO → DB insert template +
    files + extracted placeholders → return detail response.

Pipeline for ``replace_files`` (re-upload to bump version):
    delete old MinIO objects → insert new files → re-extract placeholders →
    preserve host edits where keys still match.
"""
from __future__ import annotations

import logging
import uuid

from app.core.config import settings
from app.core.storage import StorageClient, get_storage
from app.db.session import unit_of_work
from app.repositories.applicants import applicant_repo
from app.repositories.inquiries.inquiry_repo import get_by_applicant_inquiry_id
from app.repositories.leases import (
    lease_template_file_repo,
    lease_template_placeholder_repo,
    lease_template_repo,
    signed_lease_repo,
)
from app.schemas.leases.lease_template_file_response import (
    LeaseTemplateFileResponse,
)
from app.schemas.leases.lease_template_list_response import (
    LeaseTemplateListResponse,
)
from app.schemas.leases.lease_template_placeholder_response import (
    LeaseTemplatePlaceholderResponse,
)
from app.schemas.leases.lease_template_response import LeaseTemplateResponse
from app.schemas.leases.lease_template_summary import LeaseTemplateSummary
from app.services.leases.attachment_response_builder import (
    attach_presigned_urls_to_files,
)
from app.services.leases.computed import ComputedExprError, validate_expr
from app.services.leases.default_source_map import (
    get_default,
    guess_display_label,
)
from app.services.leases.default_source_resolver import (
    resolve_default_source,
    validate_default_source_spec,
)
from app.schemas.leases.suggest_placeholders_response import (
    SuggestPlaceholdersResponse,
    SuggestedPlaceholderItem,
)
from app.services.leases.placeholder_extractor import (
    extract_placeholders_across_files,
)
from app.services.leases.template_placeholder_extractor import (
    suggest_placeholders as _ai_suggest_placeholders,
)

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Allowlist for upload content types
# ---------------------------------------------------------------------------

ALLOWED_TEMPLATE_MIME_TYPES: frozenset[str] = frozenset({
    "text/markdown",
    "text/plain",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
})

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# ---------------------------------------------------------------------------
# Exceptions
# ---------------------------------------------------------------------------

class TemplateNotFoundError(LookupError):
    pass


class StorageNotConfiguredError(RuntimeError):
    pass


class TemplateFileTooLargeError(ValueError):
    pass


class TemplateFileTypeRejectedError(ValueError):
    pass


class TemplateInUseError(RuntimeError):
    """Template cannot be soft-deleted — active signed leases reference it."""


class InvalidComputedExprError(ValueError):
    pass


class InvalidDefaultSourceError(ValueError):
    pass


class PlaceholderNotFoundError(LookupError):
    pass


class ApplicantNotFoundError(LookupError):
    pass


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _normalise_content_type(declared: str | None, filename: str) -> str:
    """Resolve a content type when the upload omits one or sends octet-stream.

    DOCX uploads frequently arrive as ``application/octet-stream`` from
    browser drag-drop, so we fall back to the filename extension.
    """
    if declared and declared in ALLOWED_TEMPLATE_MIME_TYPES:
        return declared
    lower = filename.lower()
    if lower.endswith(".docx"):
        return DOCX_MIME
    if lower.endswith(".md"):
        return "text/markdown"
    if lower.endswith(".txt"):
        return "text/plain"
    if declared:
        return declared
    return "application/octet-stream"


def _extract_text_from_upload(content: bytes, content_type: str) -> str:
    """Pull plain text out of the upload for placeholder extraction.

    For DOCX we use python-docx if available; otherwise we can't extract
    placeholders (the host can still edit the spec by hand later).
    """
    if content_type in ("text/markdown", "text/plain"):
        try:
            return content.decode("utf-8", errors="replace")
        except Exception:  # noqa: BLE001
            return ""
    if content_type == DOCX_MIME:
        try:
            import io
            import docx  # type: ignore[import-untyped]

            document = docx.Document(io.BytesIO(content))
            parts = [p.text for p in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            parts.append(p.text)
            return "\n".join(parts)
        except ImportError:
            logger.info(
                "python-docx not installed — DOCX placeholder extraction skipped",
            )
            return ""
        except Exception:  # noqa: BLE001
            logger.warning("Failed to extract text from DOCX", exc_info=True)
            return ""
    return ""


def _build_summary_from_orm(template, *, file_count: int, placeholder_count: int) -> LeaseTemplateSummary:
    return LeaseTemplateSummary(
        id=template.id,
        user_id=template.user_id,
        organization_id=template.organization_id,
        name=template.name,
        description=template.description,
        version=template.version,
        file_count=file_count,
        placeholder_count=placeholder_count,
        created_at=template.created_at,
        updated_at=template.updated_at,
    )


def _build_detail_from_orm(template, files, placeholders) -> LeaseTemplateResponse:
    file_responses = [LeaseTemplateFileResponse.model_validate(f) for f in files]
    placeholder_responses = [
        LeaseTemplatePlaceholderResponse.model_validate(p) for p in placeholders
    ]
    file_responses = attach_presigned_urls_to_files(file_responses)
    return LeaseTemplateResponse(
        id=template.id,
        user_id=template.user_id,
        organization_id=template.organization_id,
        name=template.name,
        description=template.description,
        version=template.version,
        files=file_responses,
        placeholders=placeholder_responses,
        created_at=template.created_at,
        updated_at=template.updated_at,
    )


# ---------------------------------------------------------------------------
# Upload (POST)
# ---------------------------------------------------------------------------

async def upload_template(
    *,
    user_id: uuid.UUID,
    organization_id: uuid.UUID,
    name: str,
    description: str | None,
    files: list[tuple[str, bytes, str | None]],
) -> LeaseTemplateResponse:
    """Create a new template from an uploaded bundle of files.

    ``files`` is a list of ``(filename, content_bytes, declared_content_type)``.
    """
    storage = get_storage()
    if storage is None:
        raise StorageNotConfiguredError("Object storage is not configured")
    if not files:
        raise TemplateFileTypeRejectedError("At least one file is required")

    # Validate each file (size + type).
    normalised: list[tuple[str, bytes, str]] = []
    for filename, content, declared in files:
        if len(content) > settings.max_blackout_attachment_size_bytes:
            max_mb = settings.max_blackout_attachment_size_bytes // (1024 * 1024)
            raise TemplateFileTooLargeError(
                f"{filename} exceeds {max_mb}MB limit"
            )
        ct = _normalise_content_type(declared, filename)
        if ct not in ALLOWED_TEMPLATE_MIME_TYPES:
            raise TemplateFileTypeRejectedError(
                f"{filename}: unsupported file type ({ct}). "
                "Allowed: .md, .txt, .docx"
            )
        normalised.append((filename, content, ct))

    # Insert template row first so we have an ID for storage paths.
    async with unit_of_work() as db:
        template = await lease_template_repo.create(
            db,
            user_id=user_id,
            organization_id=organization_id,
            name=name,
            description=description,
        )
        template_id = template.id

    # Upload files to MinIO and write DB rows + extract placeholders.
    extracted_texts: list[str] = []
    uploaded_keys: list[str] = []
    try:
        for index, (filename, content, ct) in enumerate(normalised):
            file_id = uuid.uuid4()
            storage_key = f"lease-templates/{template_id}/{file_id}"
            storage.upload_file(storage_key, content, ct)
            uploaded_keys.append(storage_key)

            async with unit_of_work() as db:
                await lease_template_file_repo.create(
                    db,
                    template_id=template_id,
                    filename=filename,
                    storage_key=storage_key,
                    content_type=ct,
                    size_bytes=len(content),
                    display_order=index,
                )

            extracted_texts.append(_extract_text_from_upload(content, ct))

        # Extract placeholders across all files (deduped, ordered).
        keys = extract_placeholders_across_files(extracted_texts)
        async with unit_of_work() as db:
            for order, key in enumerate(keys):
                seed = get_default(key)
                await lease_template_placeholder_repo.create(
                    db,
                    template_id=template_id,
                    key=key,
                    display_label=guess_display_label(key),
                    input_type=seed.input_type,
                    required=True,
                    default_source=seed.default_source,
                    computed_expr=seed.computed_expr,
                    display_order=order,
                )
    except Exception:
        # Cleanup uploaded MinIO objects on partial failure.
        for key in uploaded_keys:
            try:
                storage.delete_file(key)
            except Exception:  # noqa: BLE001
                logger.warning("Failed to delete orphan template file %s", key)
        raise

    return await get_template(
        user_id=user_id,
        organization_id=organization_id,
        template_id=template_id,
    )


# ---------------------------------------------------------------------------
# List (GET)
# ---------------------------------------------------------------------------

async def list_templates(
    *,
    user_id: uuid.UUID,
    organization_id: uuid.UUID,
    limit: int = 50,
    offset: int = 0,
) -> LeaseTemplateListResponse:
    async with unit_of_work() as db:
        rows = await lease_template_repo.list_for_tenant(
            db,
            user_id=user_id,
            organization_id=organization_id,
            limit=limit,
            offset=offset,
        )
        total = await lease_template_repo.count_for_tenant(
            db,
            user_id=user_id,
            organization_id=organization_id,
        )
        items: list[LeaseTemplateSummary] = []
        for template in rows:
            files = await lease_template_file_repo.list_for_template(
                db, template_id=template.id,
            )
            placeholders = await lease_template_placeholder_repo.list_for_template(
                db, template_id=template.id,
            )
            items.append(
                _build_summary_from_orm(
                    template,
                    file_count=len(files),
                    placeholder_count=len(placeholders),
                )
            )
    return LeaseTemplateListResponse(
        items=items,
        total=total,
        has_more=(offset + len(items)) < total,
    )


# ---------------------------------------------------------------------------
# Get detail
# ---------------------------------------------------------------------------

async def get_template(
    *,
    user_id: uuid.UUID,
    organization_id: uuid.UUID,
    template_id: uuid.UUID,
) -> LeaseTemplateResponse:
    async with unit_of_work() as db:
        template = await lease_template_repo.get(
            db,
            template_id=template_id,
            user_id=user_id,
            organization_id=organization_id,
        )
        if template is None:
            raise TemplateNotFoundError(f"Template {template_id} not found")
        files = await lease_template_file_repo.list_for_template(
            db, template_id=template.id,
        )
        placeholders = await lease_template_placeholder_repo.list_for_template(
            db, template_id=template.id,
        )
    return _build_detail_from_orm(template, files, placeholders)


# ---------------------------------------------------------------------------
# Resolve generate-defaults for a template + applicant pair
# ---------------------------------------------------------------------------

async def generate_defaults(
    *,
    user_id: uuid.UUID,
    organization_id: uuid.UUID,
    template_id: uuid.UUID,
    applicant_id: uuid.UUID,
) -> list[dict]:
    """Return resolved default values for each placeholder in a template.

    For each placeholder with a ``default_source`` spec, evaluates the spec
    against the applicant row and (if the applicant has a linked inquiry) the
    inquiry row. Returns a list of dicts with ``key``, ``value``, and
    ``provenance`` fields.

    Placeholders without ``default_source`` are included with ``value=None``
    and ``provenance=None`` so the frontend can render all fields in one pass.
    """
    async with unit_of_work() as db:
        # Validate template belongs to caller.
        template = await lease_template_repo.get(
            db,
            template_id=template_id,
            user_id=user_id,
            organization_id=organization_id,
        )
        if template is None:
            raise TemplateNotFoundError(f"Template {template_id} not found")

        # Validate applicant belongs to caller.
        applicant = await applicant_repo.get(
            db,
            applicant_id=applicant_id,
            organization_id=organization_id,
            user_id=user_id,
        )
        if applicant is None:
            raise ApplicantNotFoundError(f"Applicant {applicant_id} not found")

        # Load linked inquiry if present.
        inquiry = None
        if applicant.inquiry_id is not None:
            inquiry = await get_by_applicant_inquiry_id(
                db,
                inquiry_id=applicant.inquiry_id,
                organization_id=organization_id,
                user_id=user_id,
            )

        placeholders = await lease_template_placeholder_repo.list_for_template(
            db, template_id=template_id,
        )

    defaults: list[dict] = []
    for p in placeholders:
        if p.input_type in ("signature", "computed"):
            continue
        if p.default_source:
            value, provenance = resolve_default_source(
                p.default_source, applicant, inquiry,
            )
        else:
            value, provenance = None, None
        defaults.append({
            "key": p.key,
            "value": value,
            "provenance": provenance,
        })

    return defaults


# ---------------------------------------------------------------------------
# Update template metadata (name / description)
# ---------------------------------------------------------------------------

async def update_template_metadata(
    *,
    user_id: uuid.UUID,
    organization_id: uuid.UUID,
    template_id: uuid.UUID,
    name: str | None,
    description: str | None,
) -> LeaseTemplateResponse:
    async with unit_of_work() as db:
        updated = await lease_template_repo.update_metadata(
            db,
            template_id=template_id,
            user_id=user_id,
            organization_id=organization_id,
            name=name,
            description=description,
        )
        if updated is None:
            raise TemplateNotFoundError(f"Template {template_id} not found")
    return await get_template(
        user_id=user_id,
        organization_id=organization_id,
        template_id=template_id,
    )


# ---------------------------------------------------------------------------
# Update placeholder spec
# ---------------------------------------------------------------------------

async def update_placeholder(
    *,
    user_id: uuid.UUID,
    organization_id: uuid.UUID,
    template_id: uuid.UUID,
    placeholder_id: uuid.UUID,
    fields: dict[str, object],
) -> LeaseTemplatePlaceholderResponse:
    # Tenant scope: confirm the template belongs to caller first.
    async with unit_of_work() as db:
        template = await lease_template_repo.get(
            db,
            template_id=template_id,
            user_id=user_id,
            organization_id=organization_id,
        )
        if template is None:
            raise TemplateNotFoundError(f"Template {template_id} not found")

        placeholder = (
            await lease_template_placeholder_repo.get_by_id_scoped_to_template(
                db,
                placeholder_id=placeholder_id,
                template_id=template_id,
            )
        )
        if placeholder is None:
            raise PlaceholderNotFoundError(
                f"Placeholder {placeholder_id} not found"
            )

        # Build the field dict, validating computed_expr against the DSL
        # allowlist before any write happens. The repository owns the actual
        # ORM mutations + flush — keeps the layered architecture clean.
        update_payload: dict[str, object] = {}
        if "computed_expr" in fields:
            expr = fields["computed_expr"]
            if expr is None or expr == "":
                update_payload["computed_expr"] = None
            else:
                try:
                    validate_expr(expr)  # type: ignore[arg-type]
                except ComputedExprError as exc:
                    raise InvalidComputedExprError(str(exc)) from exc
                update_payload["computed_expr"] = expr
        if "display_label" in fields and fields["display_label"] is not None:
            update_payload["display_label"] = fields["display_label"]
        if "input_type" in fields and fields["input_type"] is not None:
            update_payload["input_type"] = fields["input_type"]
        if "required" in fields and fields["required"] is not None:
            update_payload["required"] = bool(fields["required"])
        if "default_source" in fields:
            ds = fields["default_source"]
            if ds is not None and ds != "":
                try:
                    validate_default_source_spec(ds)  # type: ignore[arg-type]
                except ValueError as exc:
                    raise InvalidDefaultSourceError(str(exc)) from exc
            update_payload["default_source"] = ds
        if "display_order" in fields and fields["display_order"] is not None:
            update_payload["display_order"] = int(fields["display_order"])  # type: ignore[arg-type]

        updated = await lease_template_placeholder_repo.update_fields(
            db,
            placeholder=placeholder,
            fields=update_payload,
        )
        return LeaseTemplatePlaceholderResponse.model_validate(updated)


# ---------------------------------------------------------------------------
# Soft-delete
# ---------------------------------------------------------------------------

async def soft_delete_template(
    *,
    user_id: uuid.UUID,
    organization_id: uuid.UUID,
    template_id: uuid.UUID,
) -> None:
    async with unit_of_work() as db:
        template = await lease_template_repo.get(
            db,
            template_id=template_id,
            user_id=user_id,
            organization_id=organization_id,
        )
        if template is None:
            raise TemplateNotFoundError(f"Template {template_id} not found")
        in_use = await signed_lease_repo.has_active_lease_for_template(
            db, template_id=template_id,
        )
        if in_use:
            raise TemplateInUseError(
                "Cannot delete template — active leases reference it",
            )
        await lease_template_repo.soft_delete(
            db,
            template_id=template_id,
            user_id=user_id,
            organization_id=organization_id,
        )


# ---------------------------------------------------------------------------
# Replace files (re-upload — bumps version, preserves host placeholder edits)
# ---------------------------------------------------------------------------

async def replace_template_files(
    *,
    user_id: uuid.UUID,
    organization_id: uuid.UUID,
    template_id: uuid.UUID,
    files: list[tuple[str, bytes, str | None]],
) -> LeaseTemplateResponse:
    storage = get_storage()
    if storage is None:
        raise StorageNotConfiguredError("Object storage is not configured")

    async with unit_of_work() as db:
        template = await lease_template_repo.get(
            db,
            template_id=template_id,
            user_id=user_id,
            organization_id=organization_id,
        )
        if template is None:
            raise TemplateNotFoundError(f"Template {template_id} not found")

    # Validate uploads.
    normalised: list[tuple[str, bytes, str]] = []
    for filename, content, declared in files:
        if len(content) > settings.max_blackout_attachment_size_bytes:
            max_mb = settings.max_blackout_attachment_size_bytes // (1024 * 1024)
            raise TemplateFileTooLargeError(f"{filename} exceeds {max_mb}MB limit")
        ct = _normalise_content_type(declared, filename)
        if ct not in ALLOWED_TEMPLATE_MIME_TYPES:
            raise TemplateFileTypeRejectedError(
                f"{filename}: unsupported file type"
            )
        normalised.append((filename, content, ct))

    # Snapshot existing host placeholder edits so we can restore them.
    async with unit_of_work() as db:
        prior = await lease_template_placeholder_repo.list_for_template(
            db, template_id=template_id,
        )
        prior_by_key = {p.key: p for p in prior}

        # Delete old files (from DB; storage cleanup after).
        old_keys = await lease_template_file_repo.delete_all_for_template(
            db, template_id=template_id,
        )
        # Drop placeholders — re-extract below.
        await lease_template_placeholder_repo.delete_all_for_template(
            db, template_id=template_id,
        )

    # Best-effort storage cleanup of old files.
    for key in old_keys:
        try:
            storage.delete_file(key)
        except Exception:  # noqa: BLE001
            logger.warning("Failed to delete old template file %s", key)

    # Upload new files + re-extract.
    extracted_texts: list[str] = []
    new_keys: list[str] = []
    try:
        for index, (filename, content, ct) in enumerate(normalised):
            file_id = uuid.uuid4()
            storage_key = f"lease-templates/{template_id}/{file_id}"
            storage.upload_file(storage_key, content, ct)
            new_keys.append(storage_key)

            async with unit_of_work() as db:
                await lease_template_file_repo.create(
                    db,
                    template_id=template_id,
                    filename=filename,
                    storage_key=storage_key,
                    content_type=ct,
                    size_bytes=len(content),
                    display_order=index,
                )
            extracted_texts.append(_extract_text_from_upload(content, ct))

        keys = extract_placeholders_across_files(extracted_texts)
        async with unit_of_work() as db:
            for order, key in enumerate(keys):
                old = prior_by_key.get(key)
                if old is not None:
                    # Preserve host edits.
                    await lease_template_placeholder_repo.create(
                        db,
                        template_id=template_id,
                        key=key,
                        display_label=old.display_label,
                        input_type=old.input_type,
                        required=old.required,
                        default_source=old.default_source,
                        computed_expr=old.computed_expr,
                        display_order=order,
                    )
                else:
                    seed = get_default(key)
                    await lease_template_placeholder_repo.create(
                        db,
                        template_id=template_id,
                        key=key,
                        display_label=guess_display_label(key),
                        input_type=seed.input_type,
                        required=True,
                        default_source=seed.default_source,
                        computed_expr=seed.computed_expr,
                        display_order=order,
                    )

            await lease_template_repo.bump_version(
                db,
                template_id=template_id,
                user_id=user_id,
                organization_id=organization_id,
            )
    except Exception:
        for key in new_keys:
            try:
                storage.delete_file(key)
            except Exception:  # noqa: BLE001
                logger.warning("Failed to delete orphan re-uploaded file %s", key)
        raise

    return await get_template(
        user_id=user_id,
        organization_id=organization_id,
        template_id=template_id,
    )


# ---------------------------------------------------------------------------
# Helper: load a template's source files as text (used by signed-lease render)
# ---------------------------------------------------------------------------

async def load_template_source_texts(
    *,
    user_id: uuid.UUID,
    organization_id: uuid.UUID,
    template_id: uuid.UUID,
    storage: StorageClient,
) -> list[tuple[str, str, bytes]]:
    """Return ``[(filename, content_type, raw_bytes)]`` in display order."""
    async with unit_of_work() as db:
        template = await lease_template_repo.get(
            db,
            template_id=template_id,
            user_id=user_id,
            organization_id=organization_id,
        )
        if template is None:
            raise TemplateNotFoundError(f"Template {template_id} not found")
        files = await lease_template_file_repo.list_for_template(
            db, template_id=template_id,
        )
    out: list[tuple[str, str, bytes]] = []
    for f in files:
        out.append((f.filename, f.content_type, storage.download_file(f.storage_key)))
    return out


# ---------------------------------------------------------------------------
# AI placeholder suggestion (Phase 2)
# ---------------------------------------------------------------------------

async def suggest_ai_placeholders(
    *,
    user_id: uuid.UUID,
    organization_id: uuid.UUID,
    template_id: uuid.UUID,
) -> SuggestPlaceholdersResponse:
    """Run an AI pass over the template's files and propose placeholders.

    Fetches file content from MinIO, extracts text from each file, then calls
    Claude to propose a named + typed list of placeholders. The caller
    (route handler) must NOT persist the result — the host reviews the list
    and saves via the existing ``update_placeholder`` endpoint.

    Returns a :class:`SuggestPlaceholdersResponse` with ``suggestions``,
    ``truncated``, and an optional ``pages_note``.

    Raises :class:`TemplateNotFoundError` if the template doesn't belong to the
    caller. Storage errors are propagated (the route handler maps them to 503).
    """
    storage = get_storage()
    if storage is None:
        raise StorageNotConfiguredError("Object storage is not configured")

    async with unit_of_work() as db:
        template = await lease_template_repo.get(
            db,
            template_id=template_id,
            user_id=user_id,
            organization_id=organization_id,
        )
        if template is None:
            raise TemplateNotFoundError(f"Template {template_id} not found")
        files = await lease_template_file_repo.list_for_template(
            db, template_id=template_id,
        )

    # Extract text from every file.
    extracted_texts: list[str] = []
    for f in files:
        raw = storage.download_file(f.storage_key)
        extracted_texts.append(_extract_text_from_upload(raw, f.content_type))

    combined_text = "\n\n".join(t for t in extracted_texts if t.strip())
    result = await _ai_suggest_placeholders(combined_text)

    pages_note: str | None = None
    if result.truncated:
        pages_note = (
            "The document was too long to analyse in full — "
            "I read the first portion. Some placeholders near the end may be missing."
        )

    return SuggestPlaceholdersResponse(
        suggestions=[
            SuggestedPlaceholderItem(
                key=s.key,
                description=s.description,
                input_type=s.input_type,
            )
            for s in result.suggestions
        ],
        truncated=result.truncated,
        pages_note=pages_note,
    )
